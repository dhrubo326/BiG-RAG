"""
Test script for document translation using Google Gemini 2.5 Pro

This script demonstrates the translation accuracy of Gemini 2.5 Pro
by translating a Bangla document to English.

NOTE: Gemini API does NOT support .docx files directly.
Supported formats: PDF, images, audio, video, plain text

This script uses TEXT EXTRACTION approach (best option for .docx):
- Extracts text from .docx while preserving paragraph structure
- Sends extracted text to Gemini for translation
- Preserves paragraph breaks and basic structure

Usage:
    python test_gemini_translation.py
"""

import os
import sys
import time
from pathlib import Path
from dotenv import load_dotenv
import google.generativeai as genai
from docx import Document

# Load environment variables
load_dotenv()

# Configuration
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
INPUT_FILE = "Physics_text.docx"
OUTPUT_FILE = "Physics_text_translated.docx"
BANGLA_EXTRACTED_FILE = "Physics_text_extracted_bangla.txt"  # Save extracted Bangla text for verification
MODEL_NAME = "gemini-2.5-flash"  # Options: gemini-2.0-flash-exp, gemini-1.5-pro, gemini-1.5-flash

def read_docx(file_path):
    """
    Read text content from a .docx file while preserving structure.

    Args:
        file_path (str): Path to the .docx file

    Returns:
        str: Extracted text content with preserved paragraphs
    """
    try:
        print(f"[PROGRESS] Opening document: {file_path}...")
        doc = Document(file_path)

        # Get file size
        file_size = os.path.getsize(file_path)
        file_size_mb = file_size / (1024 * 1024)
        print(f"[INFO] File size: {file_size_mb:.2f} MB")

        print("[PROGRESS] Reading paragraphs from document...")
        full_text = []
        total_paragraphs = len(doc.paragraphs)

        for idx, para in enumerate(doc.paragraphs, 1):
            if para.text.strip():  # Skip empty paragraphs
                full_text.append(para.text)

            # Show progress every 100 paragraphs
            if idx % 100 == 0:
                print(f"[PROGRESS] Read {idx}/{total_paragraphs} paragraphs...")

        print(f"[OK] Finished reading {total_paragraphs} paragraphs")
        print(f"[INFO] Non-empty paragraphs: {len(full_text)}")
        print()

        return "\n\n".join(full_text)

    except Exception as e:
        print(f"[ERROR] Failed to read document: {e}")
        sys.exit(1)


def translate_with_gemini(text, api_key, model_name):
    """
    Translate text from Bangla to English using Google Gemini.

    Args:
        text (str): Text to translate
        api_key (str): Gemini API key
        model_name (str): Model name to use

    Returns:
        str: Translated text
    """
    try:
        # Configure Gemini API
        print("[PROGRESS] Configuring Gemini API...")
        genai.configure(api_key=api_key)
        print("[OK] API configured successfully")
        print()

        # Initialize model
        print(f"[PROGRESS] Initializing model: {model_name}...")
        model = genai.GenerativeModel(model_name)
        print("[OK] Model initialized")
        print()

        # Create translation prompt
        print("[PROGRESS] Preparing translation prompt...")
        prompt = f"""You are a professional translator specializing in Bangla to English translation.

Translate the following Bangla text to English. This is a physics textbook.

IMPORTANT INSTRUCTIONS:
- Translate ALL content completely
- Preserve paragraph structure and breaks
- Maintain technical terminology accuracy
- Keep the translation natural and fluent in English
- Preserve mathematical formulas and equations as-is
- Do NOT add explanations, notes, or commentary
- Do NOT summarize - translate everything
- Output ONLY the translated text, nothing else

Bangla Text:
{text}

English Translation:"""

        print(f"[INFO] Prompt created ({len(prompt):,} characters)")
        print(f"[INFO] Input text: {len(text):,} characters")
        print()

        # Generate translation
        print("[PROGRESS] Sending translation request to Gemini...")
        print("[INFO] This may take 30-90 seconds for large documents...")
        print("[INFO] Please wait...")
        print()

        start_time = time.time()

        response = model.generate_content(prompt)

        elapsed_time = time.time() - start_time

        print(f"[OK] Response received from Gemini API")
        print(f"[INFO] Translation took {elapsed_time:.1f} seconds")
        print()

        if not response or not response.text:
            raise ValueError("Empty response from Gemini API")

        translated_text = response.text.strip()
        print(f"[OK] Translation completed ({len(translated_text):,} characters)")
        print()

        return translated_text

    except Exception as e:
        print(f"[ERROR] Translation failed: {e}")
        print(f"[ERROR] Error type: {type(e).__name__}")
        sys.exit(1)


def save_extracted_bangla(bangla_text, output_file):
    """
    Save extracted Bangla text to a .txt file for verification.

    Args:
        bangla_text (str): Extracted Bangla text
        output_file (str): Output file path (.txt)
    """
    try:
        print(f"[PROGRESS] Saving extracted Bangla text to: {output_file}...")

        with open(output_file, "w", encoding="utf-8") as f:
            f.write(bangla_text)

        # Get file size
        file_size = os.path.getsize(output_file)
        file_size_kb = file_size / 1024

        print(f"[SUCCESS] Bangla text saved to: {output_file}")
        print(f"[INFO] File size: {file_size_kb:.2f} KB")
        print(f"[INFO] This file can be used to verify text extraction quality")
        print()

    except Exception as e:
        print(f"[ERROR] Failed to save Bangla text: {e}")
        sys.exit(1)


def save_translation(translated_text, output_file):
    """
    Save translated text to a .docx file.

    Args:
        translated_text (str): Translated text
        output_file (str): Output file path (.docx)
    """
    try:
        print(f"[PROGRESS] Creating Word document: {output_file}...")

        # Create a new Document
        doc = Document()

        print("[PROGRESS] Adding translated text to document...")

        # Split text into paragraphs and add to document
        paragraphs = translated_text.split('\n\n')
        total_paragraphs = len([p for p in paragraphs if p.strip()])

        print(f"[INFO] Found {total_paragraphs} paragraphs to write")

        for idx, para_text in enumerate(paragraphs, 1):
            if para_text.strip():  # Skip empty paragraphs
                doc.add_paragraph(para_text.strip())

                # Show progress every 50 paragraphs
                if idx % 50 == 0:
                    print(f"[PROGRESS] Written {idx}/{total_paragraphs} paragraphs...")

        print("[PROGRESS] Saving document to disk...")

        # Save the document
        doc.save(output_file)

        # Get file size
        file_size = os.path.getsize(output_file)
        file_size_kb = file_size / 1024

        print(f"[SUCCESS] Translation saved to: {output_file}")
        print(f"[INFO] Output file size: {file_size_kb:.2f} KB")
        print()

    except Exception as e:
        print(f"[ERROR] Failed to save translation: {e}")
        sys.exit(1)


def main():
    """Main execution function."""
    # Set UTF-8 encoding for Windows console output
    if sys.platform == 'win32':
        import codecs
        sys.stdout = codecs.getwriter('utf-8')(sys.stdout.buffer, 'strict')
        sys.stderr = codecs.getwriter('utf-8')(sys.stderr.buffer, 'strict')

    print("="*80)
    print("Gemini Translation Test - Bangla to English")
    print("="*80)
    print()
    print("[INFO] This script uses text extraction from .docx files")
    print("[INFO] Note: Gemini API does NOT support .docx file upload directly")
    print("[INFO] Supported file formats: PDF, images, audio, video, plain text")
    print()
    print("="*80)
    print()

    # Validate API key
    if not GEMINI_API_KEY:
        print("[ERROR] GEMINI_API_KEY not found in .env file")
        print("[INFO] Please add your Gemini API key to .env:")
        print("       GEMINI_API_KEY=your-api-key-here")
        sys.exit(1)

    # Check if input file exists
    if not os.path.exists(INPUT_FILE):
        print(f"[ERROR] Input file not found: {INPUT_FILE}")
        print(f"[INFO] Please ensure {INPUT_FILE} exists in the current directory")
        sys.exit(1)

    print(f"[INFO] Input file: {INPUT_FILE}")
    print(f"[INFO] Output file: {OUTPUT_FILE}")
    print(f"[INFO] Model: {MODEL_NAME}")
    print()

    # Step 1: Read document
    print("="*80)
    print("[STEP 1] Reading document")
    print("="*80)
    print()

    bangla_text = read_docx(INPUT_FILE)
    print(f"[SUCCESS] Extracted {len(bangla_text):,} characters from document")
    print()

    # Save extracted Bangla text for verification
    save_extracted_bangla(bangla_text, BANGLA_EXTRACTED_FILE)

    # Show a preview of the Bangla text
    print("[PREVIEW] First 500 characters of Bangla text:")
    print("-" * 80)
    preview_text = bangla_text[:500] + "..." if len(bangla_text) > 500 else bangla_text
    try:
        print(preview_text)
    except UnicodeEncodeError:
        print("[INFO] (Bangla text preview - contains Unicode characters)")
    print("-" * 80)
    print()

    # Step 2: Translate
    print("="*80)
    print("[STEP 2] Translating with Gemini")
    print("="*80)
    print()

    english_text = translate_with_gemini(bangla_text, GEMINI_API_KEY, MODEL_NAME)

    print(f"[SUCCESS] Translation completed: {len(english_text):,} characters")
    print()

    # Show a preview of the translated text
    print("[PREVIEW] First 800 characters of English translation:")
    print("-" * 80)
    preview_length = min(800, len(english_text))
    print(english_text[:preview_length] + ("..." if len(english_text) > 800 else ""))
    print("-" * 80)
    print()

    # Step 3: Save translation
    print("="*80)
    print("[STEP 3] Saving translation to Word document")
    print("="*80)
    print()

    save_translation(english_text, OUTPUT_FILE)

    # Summary
    print("="*80)
    print("Translation Summary")
    print("="*80)
    print(f"Input file:                {INPUT_FILE}")
    print(f"Extracted Bangla text:     {BANGLA_EXTRACTED_FILE}")
    print(f"Translated output file:    {OUTPUT_FILE}")
    print(f"Input text length:         {len(bangla_text):,} characters")
    print(f"Output text length:        {len(english_text):,} characters")
    print(f"Model used:                {MODEL_NAME}")
    print(f"Method:                    Text extraction from .docx")
    print()
    print("[DONE] Translation test completed successfully!")
    print()
    print("[TIP] Check '{BANGLA_EXTRACTED_FILE}' to verify text extraction quality")
    print("[TIP] Compare original .docx with extracted text to see if anything was lost")
    print("="*80)


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n[INFO] Translation interrupted by user")
        sys.exit(0)
    except Exception as e:
        print(f"\n[ERROR] Unexpected error: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
